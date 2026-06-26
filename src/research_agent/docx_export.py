"""Direct DOCX (Microsoft Word) export of a report.

Reuses the pure Markdown block parser from ``pdf_export`` and renders the blocks
into a ``.docx`` using ``python-docx`` (an optional dependency). Word handles
Unicode natively, so Vietnamese and other non-Latin text render correctly.

If ``python-docx`` is unavailable, ``render_docx_bytes`` raises
``DocxExportError`` and callers fall back to another export format.
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from .pdf_export import markdown_to_blocks


class DocxExportError(RuntimeError):
    """Raised when a DOCX cannot be produced (missing python-docx)."""


_HEADING_LEVEL = {"h1": 1, "h2": 2, "h3": 3}


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Append a clickable hyperlink run (blue, underlined) to a paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1A6FDB")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def render_docx_bytes(title: str, markdown: str) -> bytes:
    """Render a Markdown report to DOCX bytes.

    Raises DocxExportError if ``python-docx`` is not installed, so the caller
    can fall back to another export format.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - exercised only without python-docx
        raise DocxExportError(
            "DOCX export needs the optional 'python-docx' package. Install with: "
            "pip install \"research-agent[docx]\"."
        ) from exc

    document = Document()
    document.add_heading((title or "Research Report").strip(), level=0)

    for block in markdown_to_blocks(markdown):
        level = _HEADING_LEVEL.get(block.kind)
        if level is not None:
            document.add_heading(block.text, level=level)
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            if block.link:
                _add_hyperlink(paragraph, block.link, block.text)
            else:
                paragraph.add_run(block.text)
        elif block.link:
            paragraph = document.add_paragraph()
            _add_hyperlink(paragraph, block.link, block.text)
        else:
            document.add_paragraph(block.text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def write_docx(title: str, markdown: str, path: Path) -> Path:
    """Render and write a DOCX report to ``path``; return the written path."""
    data = render_docx_bytes(title, markdown)
    path = Path(path)
    if path.parent and not path.parent.exists():
        path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)
    return path
